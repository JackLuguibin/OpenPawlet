"""Tests for API file upload functionality (JSON base64 + multipart)."""

from __future__ import annotations

import base64
from io import BytesIO
from unittest.mock import AsyncMock, MagicMock

import httpx
import pytest

from nanobot.api.server import (
    _FileSizeExceededError,
    _parse_json_content,
    _save_base64_data_url,
    create_app,
)
from nanobot.utils.document import extract_documents


def _make_mock_agent(response_text: str = "mock response") -> MagicMock:
    agent = MagicMock()
    agent.process_direct = AsyncMock(return_value=response_text)
    agent._connect_mcp = AsyncMock()
    agent.close_mcp = AsyncMock()
    return agent


@pytest.fixture
def mock_agent():
    return _make_mock_agent()


@pytest.fixture
def app(mock_agent):
    return create_app(mock_agent, model_name="test-model", request_timeout=10.0)


# ---------------------------------------------------------------------------
# Helper function tests
# ---------------------------------------------------------------------------


def test_save_base64_data_url_saves_png(tmp_path) -> None:
    """Saving a base64 data URL creates a file with correct extension."""
    b64_data = base64.b64encode(b"fake png data").decode()
    data_url = f"data:image/png;base64,{b64_data}"
    result = _save_base64_data_url(data_url, tmp_path)
    assert result is not None
    assert result.endswith(".png")
    assert (tmp_path / result.replace(str(tmp_path) + "/", "")).read_bytes() == b"fake png data"


def test_save_base64_data_url_handles_invalid_b64(tmp_path) -> None:
    """Invalid base64 returns None."""
    result = _save_base64_data_url("data:image/png;base64,not-valid-base64!!!", tmp_path)
    assert result is None


def test_save_base64_data_url_handles_unknown_mime(tmp_path) -> None:
    """Unknown MIME type defaults to .bin."""
    b64_data = base64.b64encode(b"some data").decode()
    data_url = f"data:unknown/type;base64,{b64_data}"
    result = _save_base64_data_url(data_url, tmp_path)
    assert result is not None
    assert result.endswith(".bin")


def test_save_base64_data_url_rejects_oversized_payload(tmp_path) -> None:
    """Base64 uploads should respect the same per-file limit as multipart."""
    large_payload = base64.b64encode(b"x" * (11 * 1024 * 1024)).decode()
    data_url = f"data:image/png;base64,{large_payload}"

    with pytest.raises(_FileSizeExceededError, match="10MB limit"):
        _save_base64_data_url(data_url, tmp_path)


def test_parse_json_content_extracts_text_and_media(tmp_path) -> None:
    """Parse JSON with text + base64 image saves image and returns paths."""
    b64_data = base64.b64encode(b"img").decode()
    body = {
        "messages": [
            {
                "role": "user",
                "content": [
                    {"type": "text", "text": "describe this"},
                    {
                        "type": "image_url",
                        "image_url": {"url": f"data:image/png;base64,{b64_data}"},
                    },
                ],
            }
        ]
    }
    import os

    original_cwd = os.getcwd()
    os.chdir(tmp_path)

    try:
        text, media_paths = _parse_json_content(body)
        assert text == "describe this"
        assert len(media_paths) == 1
    finally:
        os.chdir(original_cwd)


def test_parse_json_content_plain_text_only() -> None:
    """Plain text string content returns no media."""
    body = {"messages": [{"role": "user", "content": "hello"}]}
    text, media_paths = _parse_json_content(body)
    assert text == "hello"
    assert media_paths == []


def test_parse_json_content_validates_single_message() -> None:
    """Multiple messages raise ValueError."""
    body = {
        "messages": [
            {"role": "user", "content": "first"},
            {"role": "user", "content": "second"},
        ]
    }
    with pytest.raises(ValueError, match="single user message"):
        _parse_json_content(body)


def test_parse_json_content_validates_user_role() -> None:
    """Non-user role raises ValueError."""
    body = {"messages": [{"role": "system", "content": "you are a bot"}]}
    with pytest.raises(ValueError, match="single user message"):
        _parse_json_content(body)


def test_parse_json_content_rejects_oversized_base64_file(tmp_path) -> None:
    """Oversized JSON data URLs should fail before writing to disk."""
    large_payload = base64.b64encode(b"x" * (11 * 1024 * 1024)).decode()
    body = {
        "messages": [
            {
                "role": "user",
                "content": [
                    {"type": "text", "text": "describe"},
                    {
                        "type": "image_url",
                        "image_url": {"url": f"data:image/png;base64,{large_payload}"},
                    },
                ],
            }
        ]
    }
    import os

    original_cwd = os.getcwd()
    os.chdir(tmp_path)

    try:
        with pytest.raises(_FileSizeExceededError, match="10MB limit"):
            _parse_json_content(body)
    finally:
        os.chdir(original_cwd)


# ---------------------------------------------------------------------------
# Multipart upload tests
# ---------------------------------------------------------------------------


@pytest.mark.asyncio
async def test_multipart_upload_saves_file(mock_agent, tmp_path) -> None:
    """Multipart upload saves file to media dir and passes path to process_direct."""
    import os

    original_cwd = os.getcwd()
    os.chdir(tmp_path)

    try:
        app = create_app(mock_agent, model_name="m")
        async with httpx.AsyncClient(
            transport=httpx.ASGITransport(app=app), base_url="http://testserver"
        ) as client:
            file_data = b"test file content"
            files = {"files": ("upload.bin", file_data, "application/octet-stream")}
            data = {"message": "analyze this"}
            resp = await client.post("/v1/chat/completions", data=data, files=files)
        assert resp.status_code == 200
        call_kwargs = mock_agent.process_direct.call_args.kwargs
        assert call_kwargs["content"] == "analyze this"
        assert len(call_kwargs.get("media") or []) == 1
    finally:
        os.chdir(original_cwd)


@pytest.mark.asyncio
async def test_multipart_multiple_files(mock_agent, tmp_path) -> None:
    """Multipart upload with multiple files saves all and passes paths."""
    import os

    original_cwd = os.getcwd()
    os.chdir(tmp_path)

    try:
        app = create_app(mock_agent, model_name="m")
        async with httpx.AsyncClient(
            transport=httpx.ASGITransport(app=app), base_url="http://testserver"
        ) as client:
            files = [
                ("files", ("a.txt", b"aaa", "text/plain")),
                ("files", ("b.txt", b"bbb", "text/plain")),
            ]
            resp = await client.post(
                "/v1/chat/completions", data={"message": "analyze"}, files=files
            )
        assert resp.status_code == 200
    finally:
        os.chdir(original_cwd)


@pytest.mark.asyncio
async def test_multipart_file_size_limit(mock_agent, tmp_path) -> None:
    """File exceeding MAX_FILE_SIZE returns 413."""
    import os

    original_cwd = os.getcwd()
    os.chdir(tmp_path)

    try:
        app = create_app(mock_agent, model_name="m")
        async with httpx.AsyncClient(
            transport=httpx.ASGITransport(app=app), base_url="http://testserver"
        ) as client:
            large_data = b"x" * (11 * 1024 * 1024)
            resp = await client.post(
                "/v1/chat/completions",
                data={"message": "analyze"},
                files={"files": ("big.bin", large_data, "application/octet-stream")},
            )
        assert resp.status_code == 413
    finally:
        os.chdir(original_cwd)


@pytest.mark.asyncio
async def test_multipart_defaults_text_when_missing(mock_agent, tmp_path) -> None:
    """Multipart without message field uses default text."""
    import os

    original_cwd = os.getcwd()
    os.chdir(tmp_path)

    try:
        app = create_app(mock_agent, model_name="m")
        async with httpx.AsyncClient(
            transport=httpx.ASGITransport(app=app), base_url="http://testserver"
        ) as client:
            resp = await client.post(
                "/v1/chat/completions",
                files={"files": ("f.bin", b"content", "application/octet-stream")},
            )
        assert resp.status_code == 200
        call_kwargs = mock_agent.process_direct.call_args.kwargs
        assert call_kwargs["content"] == "请分析上传的文件"
    finally:
        os.chdir(original_cwd)


@pytest.mark.asyncio
async def test_multipart_with_session_id(mock_agent, tmp_path) -> None:
    """Multipart upload with session_id uses custom session key."""
    import os

    original_cwd = os.getcwd()
    os.chdir(tmp_path)

    try:
        app = create_app(mock_agent, model_name="m")
        async with httpx.AsyncClient(
            transport=httpx.ASGITransport(app=app), base_url="http://testserver"
        ) as client:
            resp = await client.post(
                "/v1/chat/completions",
                data={"message": "hello", "session_id": "my-session"},
                files={"files": ("f.bin", b"content", "application/octet-stream")},
            )
        assert resp.status_code == 200
        call_kwargs = mock_agent.process_direct.call_args.kwargs
        assert call_kwargs["session_key"] == "api:my-session"
    finally:
        os.chdir(original_cwd)


# ---------------------------------------------------------------------------
# Backward compatibility tests
# ---------------------------------------------------------------------------


@pytest.mark.asyncio
async def test_plain_text_backward_compat(mock_agent) -> None:
    """Plain text JSON request (no media) works as before."""
    app = create_app(mock_agent, model_name="m")
    async with httpx.AsyncClient(
        transport=httpx.ASGITransport(app=app), base_url="http://testserver"
    ) as client:
        resp = await client.post(
            "/v1/chat/completions",
            json={"messages": [{"role": "user", "content": "hello world"}]},
        )
    assert resp.status_code == 200
    body = resp.json()
    assert body["choices"][0]["message"]["content"] == "mock response"
    call_kwargs = mock_agent.process_direct.call_args.kwargs
    assert call_kwargs["content"] == "hello world"
    assert call_kwargs.get("media") is None


@pytest.mark.asyncio
async def test_json_base64_image_upload(mock_agent, tmp_path) -> None:
    """JSON request with base64 data URL saves file and passes path."""
    import os

    original_cwd = os.getcwd()
    os.chdir(tmp_path)

    try:
        app = create_app(mock_agent, model_name="m")
        async with httpx.AsyncClient(
            transport=httpx.ASGITransport(app=app), base_url="http://testserver"
        ) as client:
            tiny_png_b64 = "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAADUlEQVR42mP8z8DwHwAFBQIAX8jx0gAAAABJRU5ErkJggg=="
            resp = await client.post(
                "/v1/chat/completions",
                json={
                    "messages": [
                        {
                            "role": "user",
                            "content": [
                                {"type": "text", "text": "what is this"},
                                {
                                    "type": "image_url",
                                    "image_url": {"url": f"data:image/png;base64,{tiny_png_b64}"},
                                },
                            ],
                        }
                    ]
                },
            )
        assert resp.status_code == 200
        call_kwargs = mock_agent.process_direct.call_args.kwargs
        assert call_kwargs["content"] == "what is this"
        assert len(call_kwargs.get("media", [])) == 1
    finally:
        os.chdir(original_cwd)


# ---------------------------------------------------------------------------
# extract_documents tests (now in nanobot.utils.document)
# ---------------------------------------------------------------------------


def test_extract_documents_separates_images_from_docs(tmp_path) -> None:
    """Images stay in media; document text is appended to content."""
    from docx import Document

    png = tmp_path / "chart.png"
    png.write_bytes(b"\x89PNG\r\n\x1a\n" + b"\x00" * 100)

    doc = Document()
    doc.add_paragraph("Quarterly revenue is $5M")
    docx_path = tmp_path / "report.docx"
    doc.save(docx_path)

    text, image_paths = extract_documents("summarize", [str(png), str(docx_path)])
    assert len(image_paths) == 1
    assert image_paths[0] == str(png)
    assert "Quarterly revenue" in text
    assert "summarize" in text


def test_extract_documents_skips_extraction_errors(tmp_path, monkeypatch) -> None:
    """Document extraction errors should not leak into user text."""
    bad_file = tmp_path / "broken.docx"
    bad_file.write_text("not a docx", encoding="utf-8")

    import nanobot.utils.document as _doc

    monkeypatch.setattr(
        _doc,
        "extract_text",
        lambda _path: "[error: failed to extract DOCX: boom]",
    )

    text, image_paths = extract_documents("hello", [str(bad_file)])
    assert text == "hello"
    assert image_paths == []


def test_extract_documents_images_only(tmp_path) -> None:
    """When all files are images, text is unchanged and all paths kept."""
    png = tmp_path / "a.png"
    png.write_bytes(b"\x89PNG\r\n\x1a\n" + b"\x00" * 100)
    text, image_paths = extract_documents("describe", [str(png)])
    assert text == "describe"
    assert len(image_paths) == 1


def test_extract_documents_skips_oversized_files(tmp_path) -> None:
    """Files exceeding the size limit should be silently skipped."""
    big = tmp_path / "huge.txt"
    big.write_bytes(b"x" * 200)

    text, image_paths = extract_documents("hello", [str(big)], max_file_size=100)
    assert text == "hello"
    assert image_paths == []


def test_extract_documents_does_not_read_full_file_for_mime(tmp_path) -> None:
    """MIME detection should only read header bytes, not the entire file."""
    from pathlib import Path as _Path

    big_txt = tmp_path / "big.txt"
    big_txt.write_bytes(b"hello world " * 100_000)  # ~1.2 MB

    original_read_bytes = _Path.read_bytes
    read_sizes: list[int] = []

    def _tracking_read_bytes(self):
        data = original_read_bytes(self)
        read_sizes.append(len(data))
        return data

    import unittest.mock

    with unittest.mock.patch.object(_Path, "read_bytes", _tracking_read_bytes):
        extract_documents("test", [str(big_txt)])

    # If the full file was read for MIME detection, read_sizes would
    # contain a >1MB entry.  After the fix, only a small header is read.
    assert all(size <= 4096 for size in read_sizes), (
        f"extract_documents read full file for MIME detection: sizes={read_sizes}"
    )


# ---------------------------------------------------------------------------
# DOCX upload test — API saves file, loop layer extracts text
# ---------------------------------------------------------------------------


@pytest.mark.asyncio
async def test_docx_upload_passes_media_path(tmp_path) -> None:
    """Uploaded DOCX is saved to disk and its path passed as media.
    (Text extraction happens later in AgentLoop._process_message.)"""
    agent = _make_mock_agent("report summary")
    import os

    original_cwd = os.getcwd()
    os.chdir(tmp_path)

    try:
        app = create_app(agent, model_name="m")
        async with httpx.AsyncClient(
            transport=httpx.ASGITransport(app=app), base_url="http://testserver"
        ) as client:
            from docx import Document

            doc = Document()
            doc.add_paragraph("Total revenue: $5,000,000")
            buf = BytesIO()
            doc.save(buf)

            resp = await client.post(
                "/v1/chat/completions",
                data={"message": "summarize the report"},
                files={
                    "files": (
                        "report.docx",
                        buf.getvalue(),
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    )
                },
            )
        assert resp.status_code == 200
        call_kwargs = agent.process_direct.call_args.kwargs
        assert call_kwargs["content"] == "summarize the report"
        media = call_kwargs.get("media", [])
        assert len(media) == 1
        assert "report.docx" in media[0]
    finally:
        os.chdir(original_cwd)
